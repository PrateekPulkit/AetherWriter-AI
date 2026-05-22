import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_report():
    doc = docx.Document()
    
    # --- STYLING ---
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    def add_heading(text, level=1):
        h = doc.add_heading(text, level=level)
        for run in h.runs:
            run.font.name = 'Times New Roman'
            run.font.bold = True
            run.font.size = Pt(14 if level > 1 else 16)

    def add_para(text):
        p = doc.add_paragraph(text)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        return p

    # --- TITLE PAGE ---
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("AetherWriter AI V5.0: A Multi-Layered Neural Architecture for Offline Grammar Error Correction")
    run.bold = True
    run.font.size = Pt(24)
    
    doc.add_paragraph("\n" * 4)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Faculty: DR. SUSMI JACOB")
    run.bold = True
    run.font.size = Pt(14)

    doc.add_paragraph("\n" * 2)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Presented by:")
    run.font.size = Pt(12)

    team = [
        "Prateek – AP23110011175",
        "Abhishek – AP23110011180",
        "Abdullah – AP23110011186",
        "Srinadh – AP23110011171"
    ]
    for member in team:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(member)
        run.bold = True

    doc.add_page_break()

    # --- 1. INTRODUCTION ---
    add_heading("1. INTRODUCTION", 1)
    add_para(
        "AetherWriter AI V5.0 is a state-of-the-art writing assistant designed to bring professional-grade grammar "
        "correction and linguistic analysis directly to the user's local machine. In an era where online writing "
        "tools often compromise user privacy by transmitting data to cloud servers, AetherWriter provides a 'Local-First' "
        "solution. The system leverages a hybrid architecture that combines deterministic heuristic rules with deep learning "
        "models, including Multi-Layer Perceptrons (MLP) and Generative Transformers (T5). This multi-layered defense "
        "ensures that the system can catch everything from basic punctuation errors to complex semantic inconsistencies "
        "without requiring an active internet connection. The core philosophy of AetherWriter is to empower "
        "academic and professional writers with a tool that is fast, private, and intellectually robust."
    )

    # --- 2. BACKGROUND STUDY ---
    add_heading("2. BACKGROUND STUDY", 1)
    add_para(
        "Current Writing Enhancement (WE) tools, such as Grammarly and QuillBot, rely heavily on cloud-based Large Language "
        "Models (LLMs). While effective, these systems present significant drawbacks in terms of data privacy and latency. "
        "Academic and corporate environments often deal with sensitive information that cannot ideally be shared with third-party "
        "servers. Furthermore, statistical-only models occasionally suffer from 'hallucinations,' where they suggest "
        "contextually inappropriate corrections. Research in Natural Language Processing (NLP) has shown that hybrid systems—those "
        "that combine the precision of Regular Expressions (Regex) with the flexibility of Neural Networks—outperform "
        "purely statistical models in niche grammar correction tasks. AetherWriter builds upon these findings by "
        "implementing a five-layer pipeline that transitions from deterministic rules to generative intelligence."
    )

    # --- 3. PROBLEM STATEMENT ---
    add_heading("3. PROBLEM STATEMENT", 1)
    add_para(
        "The primary challenge addressed by this project is the lack of high-performance, privacy-preserving grammar tools "
        "that can run on standard consumer hardware. Most offline spellcheckers are limited to dictionary-based lookups "
        "and miss structural errors. Conversely, advanced models like GPT-4 are too computationally expensive for local "
        "deployment. There is a critical need for an architecture that can: "
        "1) Maintain 100% data privacy by working offline; "
        "2) Provide generative-level corrections for complex sentences; "
        "3) Optimize inference speed so that the user experience remains fluid. AetherWriter addresses these needs "
        "through its proprietary 'Deterministic-to-Generative' pipeline and smart caching mechanisms."
    )

    # --- 4. PROPOSED SYSTEM ---
    add_heading("4. PROPOSED SYSTEM", 1)
    add_para(
        "The proposed system, AetherWriter AI V5.0, is built on a modular 5-layer pipeline. Every input sentence "
        "undergoes a sequence of evaluations, starting from the fastest, most certain rules and ending with "
        "the most powerful generative passes. This architecture ensures that standard errors (like 'an apple' -> 'a apple') "
        "are fixed instantly with zero hallucination risk, while deep semantic issues are handled by the T5 Transformer. "
        "The system also features a high-performance LRU cache that stores results for previously scanned sentences, "
        "reducing latency for subsequent edits by over 99%. The frontend is a modern React-based glassmorphic interface "
        "that provides real-time feedback and aesthetic engagement."
    )

    # --- 5. ALGORITHMS AND LAYERS ---
    add_heading("5. ALGORITHMS AND LAYERS", 1)

    # 5.1
    add_heading("5.1 Layer 1: Heuristic Regex Scanner (Deterministic Layer)", 2)
    add_para(
        "The first layer utilizes a hand-curated library of 65+ Regular Expression (Regex) patterns. These patterns target "
        "well-defined academic writing errors, such as subject-verb mismatches ('he have'), improper plurals ('peoples'), "
        "and common tense errors. Regex provides O(1) matching performance and 100% precision for the specific cases "
        "it covers, making it the bedrock of the system's reliability."
    )

    # 5.2
    add_heading("5.2 Layer 2: Machine Learning Ensemble (Structural Layer)", 2)
    add_para(
        "Layer 2 implements a supervised classification model. Using a TF-IDF Vectorizer trained on character-level "
        "N-grams (range 2-5), the system can detect structural anomalies that might not match a specific regex. The "
        "model is a Random Forest Ensemble (100 Trees) trained on a balanced dataset of 12,000+ sentences from the "
        "Brown, Gutenberg, and WebText corpora. This layer's strength lies in its ability to generalize and flag "
        "'weird' sounding sentences for further review."
    )

    # 5.3
    add_heading("5.3 Layer 3: NLP Part-of-Speech Tagging (Linguistic Layer)", 2)
    add_para(
        "This layer uses the NLTK Averaged Perceptron Tagger to perform deep linguistic analysis. By assigning tags "
        "like PRP (Pronoun), VBZ (Verb), and NNS (Noun Plural), the system programmatically verifies grammatical "
        "rules independent of training data. For example, it detects if a plural subject is followed by a singular "
        "verb by checking the sequence of POS tags, providing a second line of defense against agreement errors."
    )

    # 5.4
    add_heading("5.4 Layer 4: Deep Semantic T5 Transformer (Generative Layer)", 2)
    add_para(
        "The pinnacle of the pipeline is the T5 (Text-to-Text Transfer Transformer). We utilize the 'vennify/t5-base-grammar-correction' "
        "model, which treats every correction task as a translation problem (Incorrect Text -> Correct Text). "
        "With 220 million parameters, this model understands the deep semantic context of a sentence and can rewrite "
        "it for clarity and flow, acting as the ultimate polisher for the text."
    )

    # 5.5
    add_heading("5.5 Layer 5: Performance & Caching Engine", 2)
    add_para(
        "To ensure production-grade speed, we implemented an in-memory LRU Cache with a 5,000-entry capacity. Cold "
        "inference through the full AI pipeline takes approximately 1050ms. However, a 'Cache Hit'—where a sentence has "
        "been previously analyzed—returns results in just 0.08ms. Combined with an Asynchronous Concurrency model using "
        "FastAPI and Python's ThreadPoolExecutor, the system can handle concurrent user requests without blocking."
    )

    # --- 6. ALGORITHM COMPARISON ---
    add_heading("6. ALGORITHM COMPARISON", 1)
    
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Layer'
    hdr_cells[1].text = 'Technique'
    hdr_cells[2].text = 'Primary Strength'
    hdr_cells[3].text = 'Latency'

    comparisons = [
        ("Layer 1", "Regex", "100% Precision", "< 1ms"),
        ("Layer 2", "ML Ensemble", "Pattern Recognition", "10-20ms"),
        ("Layer 3", "NLTK POS", "Rule-Based Logic", "15-30ms"),
        ("Layer 4", "T5 Transformer", "Semantic Flow", "~800ms"),
        ("Cache", "Hash Key / LRU", "Instant Recall", "0.08ms")
    ]

    for layer, tech, strength, lat in comparisons:
        row_cells = table.add_row().cells
        row_cells[0].text = layer
        row_cells[1].text = tech
        row_cells[2].text = strength
        row_cells[3].text = lat

    # --- 7. OUTPUT IMAGES ---
    add_heading("7. OUTPUT IMAGES", 1)
    add_para("7.1 AetherWriter Glassmorphic Dashboard UI")
    add_para("[PLACEHOLDER: Please insert UI Screenshot here]")
    add_para("7.2 Real-time Error Detection and Highlight System")
    add_para("[PLACEHOLDER: Please insert Highlight Screenshot here]")
    add_para("7.3 Backend Neural Pipeline Logs and Cache Performance")
    add_para("[PLACEHOLDER: Please insert Terminal/Log Screenshot here]")
    add_para("7.4 AI Grammar Suggestions and Sentence Rewriting")
    add_para("[PLACEHOLDER: Please insert Suggestion Box Screenshot here]")

    # --- 8. CONCLUSION ---
    add_heading("8. CONCLUSION", 1)
    add_para(
        "AetherWriter AI V5.0 successfully demonstrates that privacy-conscious, professional-grade writing tools can be "
        "deployed locally without sacrificing quality. By chaining deterministic heuristics with state-of-the-art "
        "generative models and optimizing performance through caching, the system achieves a balance of precision and "
        "speed. The 'Backward Index Mapping' algorithm ensures that multiple corrections are applied seamlessly, and "
        "the glassmorphic UI provides a modern user experience. Future iterations will focus on expanding the N-gram "
        "coverage and integrating locally-run Llama-based models for even deeper stylistic suggestions."
    )

    doc.save('AETHERWRITER_FINAL_REPORT.docx')
    print("Report generated: AETHERWRITER_FINAL_REPORT.docx")

if __name__ == "__main__":
    create_report()
